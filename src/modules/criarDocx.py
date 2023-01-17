from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

pathDocumentSave = './tmp/Retorno.docx'

# cria um novo documento vazio
document = Document()

def adicionarTitulo(titulo):
  # adiciona um título ao documento
  titulo1 = document.add_paragraph()
  titulo1.add_run(titulo).bold = True

def adicionarConteudo(conteudo):
  # adiciona o conteúdo ao documento
  conteudo1 = document.add_paragraph()
  conteudo1.add_run(conteudo)

def adicionarCabecalho(atributo, valor):
  cabecalho1 = document.add_paragraph()
  cabecalho1.add_run(atributo).bold = True
  cabecalho1.add_run(f'{valor}.')

def configAndSave():
  # definir fonte, tamanho, alinhamento e formatação para todo o documento
  for paragraph in document.paragraphs:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
  document.styles['Normal'].font.name = 'Arial'
  document.styles['Normal'].font.size = Pt(12)
  document.sections[0].left_margin = Cm(3)
  document.sections[0].right_margin = Cm(2)
  document.sections[0].top_margin = Cm(3)
  document.sections[0].botton_margin = Cm(2)

  # Retira os erros de português
  document.correct_spelling()

  # salva o documento
  document.save(pathDocumentSave)

def deleteDocument():
  document.delete(pathDocumentSave)